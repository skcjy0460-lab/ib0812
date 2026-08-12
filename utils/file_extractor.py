# -*- coding: utf-8 -*-
"""
file_extractor.py
==================
사용자가 업로드한 심평원 급여기준 자료(PDF / HWP / HWPX / DOCX / TXT)에서
텍스트를 추출해 표준화된 결과(ExtractedDocument)로 반환한다.

설계 원칙
---------
1. 어떤 포맷이든 "실패해도 앱이 죽지 않는다" - 예외를 잡아 사용자에게
   읽기 쉬운 경고 메시지로 변환하고, 부분적으로라도 추출된 텍스트는 살린다.
2. 추출 결과는 항상 사람이 검수할 수 있도록 페이지/섹션 경계를 표시한다.
   (청구심사 자료는 한 글자 차이로 인정/불인정이 갈릴 수 있으므로
   AI에게 넘기기 전 사용자가 원문과 대조할 수 있어야 한다.)
"""
from __future__ import annotations

import os
import zipfile
from dataclasses import dataclass, field
from typing import List, Optional
from xml.etree import ElementTree as ET

from .hwp_parser import HwpParseError, parse_hwp

SUPPORTED_EXTENSIONS = (".pdf", ".hwp", ".hwpx", ".docx", ".txt")


@dataclass
class ExtractedDocument:
    filename: str
    text: str
    warnings: List[str] = field(default_factory=list)
    page_count: Optional[int] = None
    method: str = ""


def _warn(warnings: List[str], msg: str) -> None:
    warnings.append(msg)


def extract_txt(path: str, warnings: List[str]) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    _warn(warnings, "텍스트 인코딩을 자동으로 판별하지 못해 일부 문자가 깨졌을 수 있습니다.")
    with open(path, "rb") as f:
        return f.read().decode("utf-8", errors="replace")


def extract_pdf(path: str, warnings: List[str]) -> tuple[str, int]:
    try:
        import pdfplumber
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("pdfplumber 패키지가 설치되어 있지 않습니다.") from exc

    chunks: List[str] = []
    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        empty_pages = 0
        for i, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            if not page_text.strip():
                empty_pages += 1
            chunks.append(f"[--- 페이지 {i} ---]\n{page_text}")
            # 표가 있으면 텍스트만으로는 놓치는 정보가 많으므로 별도로 붙인다.
            try:
                tables = page.extract_tables()
            except Exception:  # noqa: BLE001
                tables = []
            for t_idx, table in enumerate(tables, start=1):
                rows = ["\t".join(c if c else "" for c in row) for row in table]
                chunks.append(f"[페이지 {i} - 표 {t_idx}]\n" + "\n".join(rows))
        if empty_pages:
            _warn(
                warnings,
                f"{empty_pages}개 페이지에서 텍스트를 인식하지 못했습니다 "
                "(스캔 이미지로 저장된 페이지일 수 있습니다. OCR이 필요할 수 있음).",
            )
    return "\n\n".join(chunks), page_count


def extract_docx(path: str, warnings: List[str]) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("python-docx 패키지가 설치되어 있지 않습니다.") from exc

    document = docx.Document(path)
    parts: List[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for t_idx, table in enumerate(document.tables, start=1):
        parts.append(f"[표 {t_idx}]")
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_hwp(path: str, warnings: List[str]) -> str:
    try:
        return parse_hwp_to_text(path)
    except HwpParseError as exc:
        _warn(
            warnings,
            f"HWP 정밀 추출에 실패했습니다 ({exc}). 가능하면 파일을 PDF로 변환 후 "
            "다시 업로드하거나, 텍스트를 직접 붙여넣어 주세요.",
        )
        return ""


def parse_hwp_to_text(path: str) -> str:
    paragraphs = parse_hwp(path)
    return "\n".join(p.text for p in paragraphs)


def extract_hwpx(path: str, warnings: List[str]) -> str:
    """HWPX(OOXML 유사 zip 포맷)에서 section*.xml의 텍스트 노드를 추출."""
    parts: List[str] = []
    try:
        with zipfile.ZipFile(path) as zf:
            section_files = sorted(
                n for n in zf.namelist() if n.startswith("Contents/section") and n.endswith(".xml")
            )
            if not section_files:
                _warn(warnings, "HWPX 내부에서 본문(section) XML을 찾지 못했습니다.")
            for name in section_files:
                xml_bytes = zf.read(name)
                try:
                    root = ET.fromstring(xml_bytes)
                except ET.ParseError:
                    continue
                for elem in root.iter():
                    tag = elem.tag.split("}")[-1]
                    if tag == "t" and elem.text:
                        parts.append(elem.text)
                    if tag == "p":
                        parts.append("\n")
    except zipfile.BadZipFile as exc:
        raise RuntimeError(f"HWPX(zip) 구조를 열 수 없습니다: {exc}") from exc
    return "".join(parts)


def extract_text_from_file(uploaded_path: str, original_filename: str) -> ExtractedDocument:
    """확장자에 따라 적절한 추출기로 라우팅한다."""
    warnings: List[str] = []
    ext = os.path.splitext(original_filename)[1].lower()
    page_count = None
    method = ext.lstrip(".")

    try:
        if ext == ".txt":
            text = extract_txt(uploaded_path, warnings)
        elif ext == ".pdf":
            text, page_count = extract_pdf(uploaded_path, warnings)
        elif ext == ".docx":
            text = extract_docx(uploaded_path, warnings)
        elif ext == ".hwp":
            text = extract_hwp(uploaded_path, warnings)
        elif ext == ".hwpx":
            text = extract_hwpx(uploaded_path, warnings)
        else:
            raise RuntimeError(
                f"지원하지 않는 파일 형식입니다: {ext} "
                f"(지원 형식: {', '.join(SUPPORTED_EXTENSIONS)})"
            )
    except Exception as exc:  # noqa: BLE001
        text = ""
        _warn(warnings, f"추출 중 오류가 발생했습니다: {exc}")

    text = (text or "").strip()
    if not text and not warnings:
        _warn(warnings, "문서에서 텍스트를 찾지 못했습니다.")

    return ExtractedDocument(
        filename=original_filename,
        text=text,
        warnings=warnings,
        page_count=page_count,
        method=method,
    )
